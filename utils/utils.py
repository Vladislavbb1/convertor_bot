from docx2pdf import convert
from fpdf import FPDF
import img2pdf
from docx import Document
from PIL import Image

# Конвертация файлов

def convert_txt_to_pdf(txt_file_path, pdf_file_path):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    with open(txt_file_path, 'r', encoding='utf-8') as file:
        for line in file:
            pdf.cell(0, 10, line.encode('latin-1', 'replace').decode('latin-1'), ln=True)

    pdf.output(pdf_file_path)

def convert_docx_to_pdf(docx_file_path, pdf_file_path):
    convert(docx_file_path, pdf_file_path)

def convert_docx_to_txt(docx_file_path, txt_file_path):
    doc = Document(docx_file_path)
    with open(txt_file_path, 'w', encoding='utf-8') as txt_file:
        for para in doc.paragraphs:
            txt_file.write(para.text + '\n')

def convert_txt_to_docx(txt_file_path, docx_file_path):
    doc = Document()
    with open(txt_file_path, 'r', encoding='utf-8') as file:
        for line in file:
            doc.add_paragraph(line)
    doc.save(docx_file_path)

def convert_images_to_pdf(image_files_path, pdf_file_path):
    with open(pdf_file_path, "wb") as f:
        f.write(img2pdf.convert(image_files_path))

def convert_png_to_jpeg(png_file_path, jpeg_file_path):
    # Open the PNG image file
    with Image.open(png_file_path) as img:
        # Convert the image to RGB mode before saving as JPEG
        img = img.convert("RGB")
        img.save(jpeg_file_path, "JPEG")